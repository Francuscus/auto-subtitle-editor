# Language Learning Subtitle Editor
# Version 2.0 - External Editor Workflow
# Banner Color: #00BCD4 (Cyan)

import os
import re
import tempfile
from typing import List, Tuple
from html import unescape
import zipfile

import gradio as gr
import torch
import whisperx


# -------------------------- Config --------------------------

VERSION = "2.0"
BANNER_COLOR = "#00BCD4"  # Cyan for v2.0

LANG_MAP = {
    "auto": None, "auto-detect": None, "automatic": None,
    "hungarian": "hu", "magyar": "hu", "hun": "hu",
    "spanish": "es", "español": "es", "esp": "es",
    "english": "en", "eng": "en",
}


# -------------------------- Utilities --------------------------

def normalize_lang(s: str | None):
    if not s:
        return None
    t = s.strip().lower()
    if t in LANG_MAP:
        return LANG_MAP[t]
    m = re.search(r"\b([a-z]{2,3})\b", t)
    return m.group(1) if m else None


def seconds_to_timestamp(t: float) -> str:
    """Convert seconds to SRT/ASS timestamp format"""
    t = max(t, 0.0)
    h = int(t // 3600)
    t -= h * 3600
    m = int(t // 60)
    t -= m * 60
    s = int(t)
    ms = int(round((t - s) * 1000))
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


# -------------------------- ASR Model --------------------------

_asr_model = None


def get_asr_model():
    global _asr_model
    if _asr_model is not None:
        return _asr_model

    use_cuda = torch.cuda.is_available()
    device = "cuda" if use_cuda else "cpu"
    compute = "float16" if use_cuda else "int8"
    
    print(f"[v{VERSION}] Loading WhisperX on {device} with {compute}")
    
    try:
        _asr_model = whisperx.load_model("small", device=device, compute_type=compute)
    except ValueError as e:
        if "compute type" in str(e).lower():
            fallback = "int16" if device == "cpu" else "float32"
            _asr_model = whisperx.load_model("small", device=device, compute_type=fallback)
        else:
            raise
    
    return _asr_model


# -------------------------- Transcription --------------------------

def transcribe_with_words(audio_path: str, language_code: str | None) -> List[dict]:
    """Transcribe and return word-level timestamps"""
    model = get_asr_model()
    
    print("Transcribing...")
    result = model.transcribe(audio_path, language=language_code)
    segments = result["segments"]
    
    # Extract words with timestamps
    word_segments = []
    for seg in segments:
        if "words" in seg and seg["words"]:
            for w in seg["words"]:
                word_segments.append({
                    "start": float(w.get("start", seg["start"])),
                    "end": float(w.get("end", seg["end"])),
                    "text": w.get("word", w.get("text", "")).strip(),
                })
        else:
            # Fallback: split by spaces
            text = seg["text"].strip()
            words = text.split()
            if words:
                seg_start = float(seg["start"])
                seg_end = float(seg["end"])
                seg_dur = seg_end - seg_start
                word_dur = seg_dur / len(words)
                
                for i, word in enumerate(words):
                    w_start = seg_start + (i * word_dur)
                    w_end = w_start + word_dur
                    word_segments.append({
                        "start": round(w_start, 3),
                        "end": round(w_end, 3),
                        "text": word
                    })
    
    print(f"Got {len(word_segments)} words")
    return word_segments


# -------------------------- DOCX Export for External Editing --------------------------

def export_to_docx_for_editing(word_segments: List[dict]) -> str:
    """Export words to a DOCX file that can be edited in Word/Google Docs"""
    from docx import Document
    from docx.shared import RGBColor, Pt
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Edit Your Subtitles', 0)
    
    # Instructions
    instructions = doc.add_paragraph()
    instructions.add_run('INSTRUCTIONS:\n').bold = True
    instructions.add_run(
        '1. Edit the text below (fix typos, change words)\n'
        '2. Use the color highlighting tools in Word/Google Docs to color words\n'
        '3. Save this file\n'
        '4. Upload it back to the app\n'
        '5. The app will preserve your text edits and colors!\n\n'
    )
    
    instructions.add_run('Color Guide:\n').bold = True
    instructions.add_run(
        '• Red = Highlight important words\n'
        '• Yellow = Highlight verbs\n'
        '• Cyan/Blue = Highlight nouns\n'
        '• Green = Highlight adjectives\n'
        '• Purple = Highlight endings/conjugations\n'
        '• Or use any colors you want!\n\n'
    )
    
    # Add separator
    doc.add_paragraph('─' * 80)
    
    # Add all words as a continuous paragraph
    para = doc.add_paragraph()
    for word_info in word_segments:
        word_text = word_info['text']
        run = para.add_run(word_text + ' ')
        run.font.size = Pt(14)
    
    # Save to temp file
    temp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(temp_dir, "subtitles_for_editing.docx")
    doc.save(docx_path)
    
    return docx_path


# -------------------------- Import DOCX with Colors --------------------------

def import_from_docx(docx_path: str, original_words: List[dict]) -> List[dict]:
    """Import edited DOCX and extract text + colors"""
    from docx import Document
    from docx.shared import RGBColor
    
    doc = Document(docx_path)
    
    # Extract words with their colors
    edited_words = []
    
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            
            # Split by spaces
            words = text.split()
            
            for word in words:
                if not word:
                    continue
                
                # Get color (from highlight or font color)
                color = "#FFFFFF"  # Default white
                
                # Check highlight color
                if run.font.highlight_color:
                    # Map Word highlight colors to hex
                    highlight_map = {
                        1: "#FFFF00",  # Yellow
                        2: "#00FF00",  # Bright Green
                        3: "#00FFFF",  # Turquoise/Cyan
                        4: "#FF00FF",  # Pink
                        5: "#0000FF",  # Blue
                        6: "#FF0000",  # Red
                        7: "#800000",  # Dark Red
                        9: "#008000",  # Dark Green
                        11: "#008080",  # Teal
                        12: "#000080",  # Dark Blue
                        13: "#800080",  # Violet/Purple
                    }
                    color = highlight_map.get(run.font.highlight_color, "#FFFF00")
                
                # Or check font color
                elif run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    color = f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
                
                edited_words.append({
                    "text": word,
                    "color": color
                })
    
    # Match edited words back to timestamps
    # We'll try to align by position
    result = []
    
    for i, orig_word in enumerate(original_words):
        if i < len(edited_words):
            result.append({
                "start": orig_word["start"],
                "end": orig_word["end"],
                "text": edited_words[i]["text"],
                "color": edited_words[i]["color"]
            })
        else:
            # Keep original if edited is shorter
            result.append({
                "start": orig_word["start"],
                "end": orig_word["end"],
                "text": orig_word["text"],
                "color": "#FFFFFF"
            })
    
    return result


# -------------------------- Export ASS --------------------------

def export_to_ass(words: List[dict], words_per_line: int = 5, 
                  font: str = "Arial", size: int = 36) -> str:
    """Export to ASS format with colors"""
    
    header = f"""[Script Info]
ScriptType: v4.00+
PlayResX: 1280
PlayResY: 720
ScaledBorderAndShadow: yes

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Default,{font},{size},&H00FFFFFF,&H00FFFFFF,&H00000000,&H80000000,0,0,0,0,100,100,0,0,1,2,1,2,30,30,30,1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""
    
    # Group words into subtitle lines
    events = []
    i = 0
    
    while i < len(words):
        chunk = words[i:i + words_per_line]
        if not chunk:
            break
        
        start = seconds_to_timestamp(chunk[0]["start"]).replace(",", ".")
        end = seconds_to_timestamp(chunk[-1]["end"]).replace(",", ".")
        
        # Build colored text
        text_parts = []
        for word_info in chunk:
            word_text = word_info["text"]
            color_hex = word_info.get("color", "#FFFFFF")
            
            if color_hex.startswith("#"):
                color_hex = color_hex[1:]
            
            r = int(color_hex[0:2], 16) if len(color_hex) >= 2 else 255
            g = int(color_hex[2:4], 16) if len(color_hex) >= 4 else 255
            b = int(color_hex[4:6], 16) if len(color_hex) >= 6 else 255
            
            ass_color = f"&H00{b:02X}{g:02X}{r:02X}"
            text_parts.append(f"{{\\c{ass_color}}}{word_text}")
        
        colored_text = " ".join(text_parts)
        events.append(f"Dialogue: 0,{start},{end},Default,,0,0,0,,{colored_text}")
        
        i += words_per_line
    
    return header + "\n".join(events) + "\n"


def export_to_srt(words: List[dict], words_per_line: int = 5) -> str:
    """Export to SRT format (no colors)"""
    lines = []
    i = 0
    subtitle_num = 1
    
    while i < len(words):
        chunk = words[i:i + words_per_line]
        if not chunk:
            break
        
        start = seconds_to_timestamp(chunk[0]["start"])
        end = seconds_to_timestamp(chunk[-1]["end"])
        text = " ".join(w["text"] for w in chunk)
        
        lines.append(f"{subtitle_num}\n{start} --> {end}\n{text}\n")
        subtitle_num += 1
        i += words_per_line
    
    return "\n".join(lines)


def save_file(content: str, extension: str) -> str:
    """Save to temp file"""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, f"subtitles{extension}")
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    return file_path


# -------------------------- Gradio App --------------------------

def create_app():
    with gr.Blocks(
        theme=gr.themes.Soft(),
        title="Language Learning Subtitle Editor"
    ) as app:
        
        gr.HTML(f"""
        <div style="background: {BANNER_COLOR}; color: white; padding: 20px; text-align: center; border-radius: 8px; margin-bottom: 20px;">
            <h1 style="margin: 0;">Language Learning Subtitle Editor</h1>
            <p style="margin: 5px 0 0 0;">Version {VERSION} - Edit in YOUR Word Processor!</p>
        </div>
        """)
        
        gr.Markdown("""
        ## How it works:
        1. **Transcribe** your audio
        2. **Download DOCX** file
        3. **Edit in Word/Google Docs** - Fix text, add colors with highlighter
        4. **Upload back** your edited file
        5. **Export ASS** with your colors!
        
        **No more fighting with web editors - use the tools you know!**
        """)
        
        # State
        word_segments_state = gr.State([])
        edited_words_state = gr.State([])
        
        with gr.Row():
            with gr.Column():
                gr.Markdown("### Step 1: Transcribe Audio")
                
                audio_input = gr.Audio(label="Upload Audio/Video", type="filepath")
                
                language_dropdown = gr.Dropdown(
                    choices=[
                        ("Auto-detect", "auto"),
                        ("Spanish", "es"),
                        ("Hungarian", "hu"),
                        ("English", "en")
                    ],
                    value="auto",
                    label="Language"
                )
                
                transcribe_btn = gr.Button("Transcribe", variant="primary", size="lg")
                
                status_text = gr.Textbox(
                    label="Status",
                    value="Ready to transcribe...",
                    interactive=False,
                    lines=3
                )
                
                transcript_preview = gr.Textbox(
                    label="Transcript Preview",
                    lines=10,
                    interactive=False
                )
        
        with gr.Row():
            with gr.Column():
                gr.Markdown("### Step 2: Download for Editing")
                
                download_docx_btn = gr.Button("Download DOCX for Editing", size="lg")
                docx_file = gr.File(label="Download this file, edit it in Word/Google Docs")
                
                gr.Markdown("""
                **In Word/Google Docs:**
                - Fix any transcription errors
                - Use the **Highlight** tool to color words
                - Yellow = Verbs, Red = Important, Cyan = Nouns, etc.
                - Save the file
                """)
        
        with gr.Row():
            with gr.Column():
                gr.Markdown("### Step 3: Upload Edited File")
                
                upload_docx = gr.File(label="Upload your edited DOCX", file_types=[".docx"])
                
                import_btn = gr.Button("Import Edited File", variant="primary", size="lg")
                
                import_status = gr.Textbox(
                    label="Import Status",
                    value="",
                    interactive=False,
                    lines=2
                )
        
        with gr.Row():
            with gr.Column():
                gr.Markdown("### Step 4: Export Subtitles")
                
                with gr.Row():
                    words_per_line = gr.Slider(
                        minimum=2,
                        maximum=10,
                        value=5,
                        step=1,
                        label="Words per subtitle line"
                    )
                    
                    font_family = gr.Dropdown(
                        choices=["Arial", "Times New Roman", "Courier New", "Georgia", "Verdana"],
                        value="Arial",
                        label="Font"
                    )
                    
                    font_size = gr.Slider(
                        minimum=20,
                        maximum=72,
                        value=36,
                        step=2,
                        label="Size"
                    )
                
                with gr.Row():
                    export_srt_btn = gr.Button("Export SRT (no colors)")
                    export_ass_btn = gr.Button("Export ASS (with colors)", variant="primary")
                
                srt_file = gr.File(label="SRT File")
                ass_file = gr.File(label="ASS File")
        
        # Events
        
        def do_transcribe(audio_path, language):
            if not audio_path:
                return "Error: No audio file", [], ""
            
            try:
                yield "Loading AI model...", [], ""
                
                lang_code = normalize_lang(language)
                
                yield "Transcribing... this may take a minute...", [], ""
                
                word_segments = transcribe_with_words(audio_path, lang_code)
                
                # Create preview
                preview_text = " ".join(w["text"] for w in word_segments)
                
                success = f"✅ Done! Got {len(word_segments)} words.\n\nNow click 'Download DOCX for Editing' below."
                
                return success, word_segments, preview_text
                
            except Exception as e:
                error_msg = f"❌ Error: {str(e)}"
                return error_msg, [], ""
        
        def do_download_docx(word_segments):
            if not word_segments:
                gr.Warning("Transcribe first!")
                return None
            
            try:
                docx_path = export_to_docx_for_editing(word_segments)
                return docx_path
            except Exception as e:
                gr.Warning(f"Error creating DOCX: {str(e)}")
                return None
        
        def do_import_docx(uploaded_file, original_words):
            if not uploaded_file:
                return [], "No file uploaded"
            
            if not original_words:
                return [], "Transcribe first!"
            
            try:
                edited_words = import_from_docx(uploaded_file.name, original_words)
                preview = " ".join(w["text"] for w in edited_words[:50])
                if len(edited_words) > 50:
                    preview += "..."
                
                return edited_words, f"✅ Imported {len(edited_words)} words!\n\nPreview: {preview}"
            except Exception as e:
                return [], f"❌ Error importing: {str(e)}"
        
        def do_export_srt(words, words_per):
            if not words:
                gr.Warning("Import edited file first!")
                return None
            
            srt_content = export_to_srt(words, int(words_per))
            return save_file(srt_content, ".srt")
        
        def do_export_ass(words, words_per, font, size):
            if not words:
                gr.Warning("Import edited file first!")
                return None
            
            ass_content = export_to_ass(words, int(words_per), font, int(size))
            return save_file(ass_content, ".ass")
        
        # Connect events
        transcribe_btn.click(
            fn=do_transcribe,
            inputs=[audio_input, language_dropdown],
            outputs=[status_text, word_segments_state, transcript_preview]
        )
        
        download_docx_btn.click(
            fn=do_download_docx,
            inputs=[word_segments_state],
            outputs=[docx_file]
        )
        
        import_btn.click(
            fn=do_import_docx,
            inputs=[upload_docx, word_segments_state],
            outputs=[edited_words_state, import_status]
        )
        
        export_srt_btn.click(
            fn=do_export_srt,
            inputs=[edited_words_state, words_per_line],
            outputs=[srt_file]
        )
        
        export_ass_btn.click(
            fn=do_export_ass,
            inputs=[edited_words_state, words_per_line, font_family, font_size],
            outputs=[ass_file]
        )
    
    return app


if __name__ == "__main__":
    app = create_app()
    app.launch()
